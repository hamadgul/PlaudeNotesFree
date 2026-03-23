#!/usr/bin/env python3
from __future__ import annotations
"""
Plaud → Whisper → Summary → DOCX Pipeline
==========================================
Drop Plaud MP3 exports into the `input/` folder.
Transcribes with local Whisper (FREE), summarizes with Claude API,
saves a formatted .docx to `output/`.

Setup:
    pip install openai-whisper anthropic python-docx watchdog
    pip install torch torchvision torchaudio  # for Whisper

Usage:
    python plaud_summarizer.py              # watch mode (auto-processes new files)
    python plaud_summarizer.py file.mp3     # process a single file
    python plaud_summarizer.py --model base --output /path/to/out file.mp3
"""

import json
import os
import subprocess
import sys
import time
import argparse
import datetime
import traceback
from pathlib import Path

# ── Config ──────────────────────────────────────────────────────────────────
INPUT_DIR = Path.home() / "Library/Mobile Documents/com~apple~CloudDocs/PlaudInput"
OUTPUT_DIR = Path.home() / "Library/Mobile Documents/com~apple~CloudDocs/PlaudOutput"

WHISPER_MODEL = "small"        # Options: tiny, base, small, medium, large
                               # 'small' = best balance of speed/accuracy for meetings

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL = "claude-sonnet-4-6"  # reliable JSON output for meeting summaries

SUMMARY_PROMPT = """You are an expert meeting notes assistant. Analyze this work meeting transcript and return a JSON object with EXACTLY this structure (no markdown, no code fences, raw JSON only):

{{
  "attendees": ["Alice", "Bob"],
  "tldr": "2-3 sentence executive summary of the meeting",
  "decisions": [
    "Decision made (be specific)"
  ],
  "action_items": [
    {{"task": "What needs to be done", "owner": "Person responsible or Unknown", "deadline": "Date/timeframe or empty string"}}
  ],
  "topics": [
    {{
      "title": "Topic or agenda item",
      "points": ["Key point", "Key point"]
    }}
  ],
  "open_questions": [
    "Unresolved question or follow-up item"
  ]
}}

Rules:
- Be thorough. Capture every decision, action item, and key point.
- For action items: infer ownership from context (e.g. "John will...") even if not explicitly stated.
- If a field has no content, use an empty list [] or empty string "".
- topics should be organized by subject discussed, not by speaker.
- Return ONLY the JSON object. No explanation, no markdown.

Transcript:
{transcript}"""

# ─────────────────────────────────────────────────────────────────────────────

def load_whisper_model(model_name: str):
    """Load and return a Whisper model, auto-detecting the best available device."""
    try:
        import whisper
        import torch
    except ImportError:
        print("Whisper not installed. Run: pip install openai-whisper torch")
        sys.exit(1)

    if torch.backends.mps.is_available():
        device = "mps"
    elif torch.cuda.is_available():
        device = "cuda"
    else:
        device = "cpu"

    print(f"    Loading Whisper '{model_name}' model on {device}...")
    try:
        return whisper.load_model(model_name, device=device)
    except (NotImplementedError, RuntimeError):
        if device != "cpu":
            print(f"    {device.upper()} not fully supported, falling back to CPU...")
            return whisper.load_model(model_name, device="cpu")
        raise


def transcribe_audio(audio_path: Path, model) -> str:
    """Transcribe audio using a pre-loaded Whisper model."""
    print(f"    Transcribing {audio_path.name}...")
    result = model.transcribe(str(audio_path), verbose=False, fp16=False)
    transcript = result["text"].strip()
    print(f"    Transcribed ({len(transcript.split())} words)")
    return transcript


def summarize_with_claude(transcript: str) -> dict | None:
    """Summarize transcript using Claude API. Returns structured dict or None."""
    if not ANTHROPIC_API_KEY:
        print("   No ANTHROPIC_API_KEY found. Skipping AI summary.")
        return None

    key = ANTHROPIC_API_KEY.strip()
    try:
        key.encode('ascii')
    except UnicodeEncodeError:
        print("   ANTHROPIC_API_KEY contains non-ASCII characters. Re-set it.")
        return None

    try:
        import anthropic
    except ImportError:
        print("Anthropic SDK not installed. Run: pip install anthropic")
        return None

    client = anthropic.Anthropic(api_key=key)
    print(f"    Summarizing with Claude ({CLAUDE_MODEL})...")

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=2000,
        messages=[{
            "role": "user",
            "content": SUMMARY_PROMPT.format(transcript=transcript)
        }]
    )

    raw = message.content[0].text.strip()
    # Strip markdown code fences Claude sometimes adds despite instructions
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1]  # drop first line (```json or ```)
        raw = raw.rsplit("```", 1)[0].strip()  # drop closing ```

    try:
        data = json.loads(raw)
        print("    Summary generated")
        return data
    except json.JSONDecodeError:
        print("    Warning: Claude did not return valid JSON. Saving raw text as tldr.")
        return {"attendees": [], "tldr": raw, "decisions": [], "action_items": [], "topics": [], "open_questions": []}


def save_docx(audio_path: Path, summary: dict | None, output_dir: Path):
    """Save a formatted Word document from structured meeting summary."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        return None

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Title ─────────────────────────────────────────────────────────────────
    mtime = audio_path.stat().st_mtime
    recording_date = datetime.datetime.fromtimestamp(mtime).strftime("%B %d, %Y")
    stem = audio_path.stem.replace("_", " ").replace("-", " ").title()

    title = doc.add_heading(stem or "Meeting Notes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    meta = doc.add_paragraph()
    meta.add_run(f"{recording_date}").font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    if not summary:
        doc.add_paragraph("No AI summary generated (set ANTHROPIC_API_KEY to enable).")
    else:
        # ── Attendees ─────────────────────────────────────────────────────────
        attendees = summary.get("attendees", [])
        if attendees:
            doc.add_heading("Attendees", 1)
            doc.add_paragraph(", ".join(attendees))
            doc.add_paragraph()

        # ── TL;DR ─────────────────────────────────────────────────────────────
        tldr = summary.get("tldr", "")
        if tldr:
            doc.add_heading("Summary", 1)
            doc.add_paragraph(tldr)
            doc.add_paragraph()

        # ── Key Decisions ─────────────────────────────────────────────────────
        decisions = summary.get("decisions", [])
        if decisions:
            doc.add_heading("Key Decisions", 1)
            for d in decisions:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(d)
            doc.add_paragraph()

        # ── Action Items (table) ───────────────────────────────────────────────
        action_items = summary.get("action_items", [])
        if action_items:
            doc.add_heading("Action Items", 1)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, label in enumerate(["Task", "Owner", "Deadline"]):
                hdr[i].text = label
                for para in hdr[i].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            # Set column widths
            col_widths = [Inches(3.2), Inches(1.5), Inches(1.3)]
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = width

            for item in action_items:
                row = table.add_row().cells
                row[0].text = item.get("task", "")
                row[1].text = item.get("owner", "Unknown")
                row[2].text = item.get("deadline", "") or "—"
                for cell in row:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)
            doc.add_paragraph()

        # ── Discussion Topics ──────────────────────────────────────────────────
        topics = summary.get("topics", [])
        if topics:
            doc.add_heading("Discussion Topics", 1)
            for topic in topics:
                doc.add_heading(topic.get("title", ""), 2)
                for point in topic.get("points", []):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(point)
            doc.add_paragraph()

        # ── Open Questions ─────────────────────────────────────────────────────
        open_questions = summary.get("open_questions", [])
        if open_questions:
            doc.add_heading("Open Questions", 1)
            for q in open_questions:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(q)
            doc.add_paragraph()

    # ── Date Footer ───────────────────────────────────────────────────────────
    footer_text = doc.add_paragraph(recording_date)
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_text.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Save ──────────────────────────────────────────────────────────────────
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"{audio_path.stem}_{timestamp}.docx"
    out_path = output_dir / out_name
    doc.save(out_path)
    print(f"    Saved  {out_path}")
    return out_path


CONFIG_PATH = Path.home() / ".config" / "plaud_notes" / "config.json"


def load_config() -> dict:
    """Load persisted folder config, or return empty dict if not set."""
    if CONFIG_PATH.exists():
        try:
            return json.loads(CONFIG_PATH.read_text())
        except Exception:
            pass
    return {}


def save_config(config: dict):
    CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    CONFIG_PATH.write_text(json.dumps(config, indent=2))


def pick_folder(prompt: str) -> Path | None:
    """Show native macOS folder picker. Returns chosen Path or None on cancel."""
    import platform
    if platform.system() != "Darwin":
        return None
    script = f'POSIX path of (choose folder with prompt "{prompt}")'
    try:
        result = subprocess.run(["osascript", "-e", script],
                                capture_output=True, text=True, timeout=60)
        chosen = result.stdout.strip()
        if chosen:
            return Path(chosen)
    except Exception:
        pass
    return None


def setup_folders(config: dict) -> dict:
    """Run folder pickers for input and output, save to config."""
    print("  Setting up folders...")
    input_dir = pick_folder("Select your Plaud INPUT folder (where recordings are dropped):")
    if not input_dir:
        print("  No input folder selected, keeping default.")
        input_dir = Path(config.get("input_dir", str(INPUT_DIR)))

    output_dir = pick_folder("Select your OUTPUT folder (where notes are saved by default):")
    if not output_dir:
        print("  No output folder selected, keeping default.")
        output_dir = Path(config.get("output_dir", str(OUTPUT_DIR)))

    config["input_dir"] = str(input_dir)
    config["output_dir"] = str(output_dir)
    save_config(config)
    print(f"  Input:  {input_dir}")
    print(f"  Output: {output_dir}")
    print(f"  Saved to {CONFIG_PATH}\n")
    return config


def ask_save_location(filename: str, default_dir: Path) -> Path:
    """Show native macOS folder picker. Returns chosen folder or default_dir on cancel."""
    import platform
    if platform.system() != "Darwin":
        return default_dir

    script = f'POSIX path of (choose folder with prompt "Save \\"{filename}\\" to:")'
    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=60
        )
        chosen = result.stdout.strip()
        if chosen:
            return Path(chosen)
    except Exception:
        pass
    return default_dir


def process_file(audio_path: Path, model, output_dir: Path):
    """Full pipeline: audio → transcript → summary → docx."""
    print(f"\n{'='*55}")
    print(f"  Processing: {audio_path.name}")
    print(f"{'='*55}")

    transcript = transcribe_audio(audio_path, model)
    summary = summarize_with_claude(transcript)

    save_dir = ask_save_location(audio_path.name, output_dir)
    save_dir.mkdir(parents=True, exist_ok=True)
    out_path = save_docx(audio_path, summary, save_dir)

    if out_path:
        print(f"\n  Done!  {out_path}\n")
    return out_path


def watch_mode(whisper_model_name: str, input_dir: Path, output_dir: Path):
    """Poll input folder every 5 seconds — reliable for iCloud Drive."""
    seen = set()

    if input_dir.exists():
        for f in input_dir.iterdir():
            if f.suffix.lower() in (".mp3", ".wav", ".m4a", ".ogg"):
                seen.add(f.name)

    print(f"  Watching: {input_dir}")
    print(f"    Output:   {output_dir}")
    print(f"    Polling every 5 seconds.")
    print(f"    {len(seen)} existing file(s) skipped on startup.")
    print(f"    Press Ctrl+C to stop.\n")

    # Load model once, reuse for every file
    model = load_whisper_model(whisper_model_name)

    while True:
        try:
            if not input_dir.exists():
                print(f"  Input folder missing: {input_dir}")
                time.sleep(10)
                continue

            for f in sorted(input_dir.iterdir()):
                if f.suffix.lower() not in (".mp3", ".wav", ".m4a", ".ogg"):
                    continue
                if f.name in seen:
                    continue
                try:
                    if f.stat().st_size == 0:
                        print(f"    {f.name} still downloading, waiting...")
                        continue
                except Exception:
                    continue

                seen.add(f.name)
                try:
                    process_file(f, model, output_dir)
                except Exception as e:
                    print(f"  Error processing {f.name}: {e}")
                    traceback.print_exc()

        except KeyboardInterrupt:
            print("\nStopped.")
            sys.exit(0)
        except Exception as e:
            print(f"  Watcher error: {e}")

        time.sleep(5)


def main():
    import io
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace', line_buffering=True)
    except AttributeError:
        pass  # stdout/stderr already wrapped or no buffer (e.g. Automator)

    parser = argparse.ArgumentParser(description="Plaud → Whisper → Summary → DOCX")
    parser.add_argument("file", nargs="?", help="Audio file to process (optional; omit for watch mode)")
    parser.add_argument("--model", default=WHISPER_MODEL, help="Whisper model size (tiny/base/small/medium/large)")
    parser.add_argument("--output", default=None, help="Override output directory for this run")
    parser.add_argument("--setup", action="store_true", help="Re-run folder setup")
    args = parser.parse_args()

    config = load_config()

    # First run or --setup flag: prompt for folders
    if args.setup or "input_dir" not in config:
        config = setup_folders(config)

    input_dir = Path(config["input_dir"])
    input_dir.mkdir(parents=True, exist_ok=True)

    output_dir = Path(args.output) if args.output else Path(config["output_dir"])
    output_dir.mkdir(parents=True, exist_ok=True)

    if args.file:
        model = load_whisper_model(args.model)
        process_file(Path(args.file), model, output_dir)
    else:
        watch_mode(args.model, input_dir, output_dir)


if __name__ == "__main__":
    main()
